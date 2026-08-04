import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()
    
    # Title
    titulo = doc.add_heading('Proyecto Integrador: ADDJ MOTORS', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("Fecha de inicio: 1 de Julio de 2026\nFecha actual: 4 de Agosto de 2026")
    doc.add_paragraph("Materia: Proyecto Integrador; Grupo TI - 3D. Mayo - Agosto 2026")
    
    # --- INGLÉS ---
    doc.add_heading('1. Inglés (English Report)', level=1)
    doc.add_paragraph("The development team is working on the ADDJ Motors system. We are developing a web platform for buying and selling vehicles. The database administrator is designing the ER model and the SQL queries. The frontend developer is building the user interface.")
    doc.add_paragraph("For the next steps, we are going to deploy the application on a local server. The team is going to test all the CRUD functionalities. The project manager is going to present the final software.")
    doc.add_paragraph("Regarding team rules and tasks: The frontend developer must design responsive interfaces. The backend developer has to implement the REST API securely. All members should write clean code and communicate daily. You must not skip the QA testing phase.")
    
    # --- BASE DE DATOS Y POO ---
    doc.add_heading('2. Base de Datos y Programación Orientada a Objetos', level=1)
    
    doc.add_heading('Planteamiento del problema', level=2)
    doc.add_paragraph("Una agencia de autos desea implementar un sistema para la compra-venta de autos usados. Para esto requiere llevar un control preciso de los clientes registrados (vendedores y compradores) y el inventario de vehículos. Se requiere que el sistema permita registrar a una persona con sus datos personales y luego permitirle ofertar un vehículo con todas sus especificaciones técnicas y legales. Además, se debe facilitar la búsqueda de vehículos por diferentes criterios y generar un acta de compraventa en el momento de la transacción.")
    
    doc.add_heading('Requerimientos funcionales y no funcionales', level=2)
    doc.add_paragraph("Requerimientos Funcionales:\n- Registro y gestión de clientes (Nombre completo, Domicilio, Correo electrónico, Teléfono).\n- Registro y gestión de vehículos vinculados a un cliente vendedor (Número de motor, serie, modelo, marca, línea, color, precio de compra y venta, transmisión, cilindros, nacionalidad, descripción, observaciones).\n- Búsqueda avanzada de vehículos por modelo, marca, precio y fecha de publicación.\n- Generación de acta de compraventa con datos del vehículo, vendedor y comprador.\n- Listado de ofertas activas (excluyendo vendidos) ordenadas por fecha.\n- Listado de vehículos vendidos con su fecha de venta.\n- Registro de abonos por ventas y cancelación de ventas.")
    doc.add_paragraph("Requerimientos No Funcionales:\n- El sistema debe ser una aplicación Web accesible desde el navegador.\n- La base de datos debe ser relacional (MySQL 8).\n- El backend debe estar desarrollado en Node.js y el generador de actas en Java (POO).\n- Tiempos de respuesta óptimos para las consultas de búsqueda.")
    
    doc.add_heading('Modelos de Base de Datos y Clases', level=2)
    doc.add_paragraph("Modelo Entidad-Relación y Relacional:")
    doc.add_paragraph("Entidades principales: Clientes, Vehiculos, Ventas, Abonos_Venta, Historial_Estados.")
    doc.add_paragraph("Relaciones: Un cliente puede tener muchos vehículos (1:N). Un vehículo puede estar en una venta (1:1). Un cliente puede ser comprador en muchas ventas (1:N). Una venta puede tener múltiples abonos (1:N).")
    doc.add_paragraph("Modelo de Clases (POO): Implementado en Java (Cliente.java, Vehiculo.java, GeneradorActa.java) para abstraer las entidades y encapsular la lógica del acta HTML.")
    
    doc.add_heading('Consultas SQL (LDD y LMD)', level=2)
    doc.add_paragraph("A continuación se presenta el extracto de código SQL utilizado para la creación (LDD) y manipulación (LMD):")
    
    sql_code = """-- Consultas LDD (Lenguaje de Definición de Datos)
CREATE TABLE IF NOT EXISTS clientes (
  id_cliente INT AUTO_INCREMENT PRIMARY KEY,
  nombre_completo VARCHAR(150) NOT NULL,
  domicilio VARCHAR(255) NOT NULL,
  correo_electronico VARCHAR(120) NOT NULL UNIQUE,
  telefono VARCHAR(20) NOT NULL
);

CREATE TABLE IF NOT EXISTS vehiculos (
  id_vehiculo INT AUTO_INCREMENT PRIMARY KEY,
  id_vendedor INT NOT NULL,
  numero_motor VARCHAR(60) NOT NULL UNIQUE,
  numero_serie VARCHAR(60) NOT NULL UNIQUE,
  modelo SMALLINT NOT NULL,
  marca VARCHAR(60) NOT NULL,
  precio_venta DECIMAL(12,2) NOT NULL,
  estado ENUM('PUBLICADO','APARTADO','VENDIDO') NOT NULL DEFAULT 'PUBLICADO',
  CONSTRAINT fk_vehiculo_vendedor FOREIGN KEY (id_vendedor) REFERENCES clientes(id_cliente)
);

-- Consultas LMD (Lenguaje de Manipulación de Datos) - CRUD
-- CREATE
INSERT INTO clientes (nombre_completo, domicilio, correo_electronico, telefono) 
VALUES ('Juan Perez', 'Calle Falsa 123', 'juan@ejemplo.com', '555-1234');

-- READ
SELECT * FROM vehiculos WHERE estado = 'PUBLICADO' ORDER BY fecha_publicacion DESC;

-- UPDATE
UPDATE vehiculos SET estado = 'VENDIDO' WHERE id_vehiculo = 1;

-- DELETE
DELETE FROM vehiculos WHERE id_vehiculo = 5;"""
    doc.add_paragraph(sql_code, style='Intense Quote')
    
    doc.add_heading('Capturas de pantalla del sistema', level=2)
    doc.add_paragraph("[NOTA PARA EL ALUMNO: Inserte aquí las capturas de pantalla del sistema funcionando: Interfaz gráfica, reportes, CRUD y base de datos.]")
    
    # --- DESARROLLO DEL PENSAMIENTO ---
    doc.add_heading('3. Desarrollo del Pensamiento y Toma de Decisiones', level=1)
    doc.add_paragraph("Rúbrica a ser evaluada presencialmente por el docente:")
    doc.add_paragraph("- Formalidad\n- Facilidad de Expresión Verbal\n- Expresión Corporal\n- Dicción\n- Dominio del Tema\n- Gestión del Tiempo\n- Orden de la Presentación\n- Participación Activa\n- Respuestas del Alumno")
    
    # --- PROYECTO INTEGRADOR ---
    doc.add_heading('4. Proyecto Integrador', level=1)
    
    doc.add_heading('4.1 Nombre y objetivo del proyecto', level=2)
    doc.add_paragraph("Nombre: ADDJ MOTORS - Sistema de compra-venta de vehículos.")
    doc.add_paragraph("Objetivo: Desarrollar e implementar un sistema integral para administrar eficientemente el proceso de compra-venta de vehículos usados de una agencia, controlando clientes, vehículos, publicaciones, ventas y generación de reportes automáticos.")
    
    doc.add_heading('4.2 Diagramas (Tiempo/Esfuerzo y Gantt)', level=2)
    doc.add_paragraph("Etapas y Actividades:\n1. Análisis y Recolección de Requerimientos (1-10 Julio)\n2. Diseño de Base de Datos y Prototipos UI (11-20 Julio)\n3. Desarrollo Backend Node.js, Java y Frontend (21-31 Julio)\n4. Pruebas, Validación y Entrega Final (1-4 Agosto)")
    doc.add_paragraph("[NOTA PARA EL ALUMNO: Inserte aquí las imágenes de su Diagrama de Tiempo/Esfuerzo (Ejes X e Y) y el Diagrama de Gantt]")
    
    doc.add_heading('4.3 Determinación de Costos, Gastos y Precio de Venta', level=2)
    doc.add_paragraph("Costos Totales (CT): $25,000 MXN (Desarrollo, licencias, equipo)\nGastos Totales (GT): $5,000 MXN (Servicios, internet, administración)\nUtilidad Esperada: 25%\nImpuestos (IVA): 16%")
    doc.add_paragraph("Cálculo del Precio de Venta (PV = CT + GT + Utilidad(25%) + Impuestos(16%)):\nSubtotal = (25,000 + 5,000) = $30,000\nSubtotal con Utilidad = 30,000 * 1.25 = $37,500\nPrecio de Venta Final (con IVA) = 37,500 * 1.16 = $43,500 MXN")
    
    doc.add_heading('4.4 Riesgos en el Proyecto y Plan de Acción', level=2)
    doc.add_paragraph("Se han identificado 8 riesgos (2 por cada etapa), cuyo valor de contingencia ha sido considerado en los Costos Totales.")
    doc.add_paragraph("Etapa 1: Análisis\n- Riesgo 1: Requerimientos ambiguos. Plan de Acción: Realizar validación constante y firma de acuerdos con el cliente.\n- Riesgo 2: Cambios de alcance imprevistos. Plan de Acción: Congelar el alcance una vez firmado y cotizar cambios por separado.")
    doc.add_paragraph("Etapa 2: Diseño\n- Riesgo 3: Diseño de BD inconsistente. Plan de Acción: Normalización estricta y revisión por pares.\n- Riesgo 4: Interfaz poco intuitiva. Plan de Acción: Uso de prototipado rápido y feedback de usuarios prueba.")
    doc.add_paragraph("Etapa 3: Desarrollo\n- Riesgo 5: Retrasos en el desarrollo de módulos. Plan de Acción: Uso de metodologías ágiles (Scrum) y entregas semanales.\n- Riesgo 6: Incompatibilidad entre módulos (Java y Node). Plan de Acción: Definición clara de APIs y formatos de intercambio (JSON).")
    doc.add_paragraph("Etapa 4: Pruebas\n- Riesgo 7: Fallos no detectados. Plan de Acción: Implementar checklist de pruebas exhaustivo (Pruebas manuales).\n- Riesgo 8: Pérdida de datos en pruebas. Plan de Acción: Ejecutar scripts de reseteo y respaldos regulares.")
    
    doc.save('Proyecto_Integrador_Vehiculos.docx')
    print("Document created successfully.")

if __name__ == '__main__':
    main()
